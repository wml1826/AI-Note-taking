"""
智能笔记软件 - FastAPI 后端
代理云端大模型（DeepSeek / 通义千问），提供笔记场景接口：
  POST /api/summary          智能摘要
  POST /api/polish           学术写作辅助
  POST /api/summary/stream   流式摘要
  POST /api/polish/stream    流式润色
  POST /api/rag/index        增量建立或更新文献索引（本地 Jina + Milvus Lite）
  POST /api/rag/retrieve     混合检索并返回来源元数据（Dense+Sparse, WeightedRanker）
  POST /api/rag/delete       删除指定文献索引
  POST /api/import           导入 PDF/Word/TXT/Markdown，提取纯文本
  GET  /health               健康检查

设计说明：
- 后端持有默认配置（见 .env），也可由请求体覆盖 api_key / base_url / model
- 使用 OpenAI 兼容协议调用云端 Chat Completions 接口
- CORS 放开，允许本地 Qt 客户端访问
- 检索：本地 Jina Embedding（768 维）+ 中英文 n-gram BM25 + Milvus Lite 原生混合检索
- RAG 流程：Python 负责「分块 + 向量化 + 混合检索」，最终 Prompt 组装与 LLM 生成由 C++ 客户端完成
"""
import hashlib
import json
import math
import os
import re
from typing import Any, Dict, List, Optional

import httpx
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from pydantic import Field

from local_rag_engine import LocalRagEngine

load_dotenv()

# ---- 本地 RAG 引擎（惰性单例：首次调用时才加载 Jina 模型与 Milvus）----
_RAG_ENGINE: Optional[LocalRagEngine] = None


def _engine() -> LocalRagEngine:
    """惰性获取本地 RAG 引擎单例。"""
    global _RAG_ENGINE
    if _RAG_ENGINE is None:
        base_dir = os.getenv("LOCAL_RAG_DIR", ".")
        _RAG_ENGINE = LocalRagEngine(base_dir=base_dir)
    return _RAG_ENGINE


app = FastAPI(title="智能笔记软件 API", version="1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---- 默认服务商配置（可被请求体覆盖）----
DEFAULT_PROVIDERS = {
    "deepseek": {
        "api_key": os.getenv("DEEPSEEK_API_KEY", ""),
        "base_url": os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com/v1"),
        "model": os.getenv("DEEPSEEK_MODEL", "deepseek-chat"),
    },
    "dashscope": {
        "api_key": os.getenv("DASHSCOPE_API_KEY", ""),
        "base_url": os.getenv("DASHSCOPE_BASE_URL",
                              "https://dashscope.aliyuncs.com/compatible-mode/v1"),
        "model": os.getenv("DASHSCOPE_MODEL", "qwen-plus"),
    },
}


# ---- 请求模型 ----
class TextRequest(BaseModel):
    text: str = Field(..., max_length=200000)
    provider: str = "deepseek"          # deepseek / dashscope
    detail: str = "medium"              # short / medium / detailed（仅摘要用）
    api_key: Optional[str] = None       # 可选：覆盖默认
    base_url: Optional[str] = None
    model: Optional[str] = None


# ---- RAG 请求模型 ----
class RagIndexRequest(BaseModel):
    note_id: str                        # 文献唯一 id（C++ 端生成）
    text: str = Field(..., max_length=1000000)
    title: str = ""
    source: str = ""


class RagRetrieveRequest(BaseModel):
    query: str = Field(..., max_length=5000)
    # None 或空列表表示主动检索全部文献；非空则限定在这些文献内
    document_ids: Optional[List[str]] = None
    top_k: int = 10
    history: List[Dict[str, Any]] = Field(default_factory=list)  # [{"role","content"}]
    rewrite_provider: Optional[str] = None
    rewrite_api_key: Optional[str] = None
    rewrite_base_url: Optional[str] = None
    rewrite_model: Optional[str] = None


class RagDeleteRequest(BaseModel):
    note_id: str


# ====================================================================
# 文本处理工具：分词 / BM25 / 页码 / 语义分块 / 哈希 / RRF / 查询改写
# （均被 test_rag_engine.py 覆盖，修改需同步测试）
# ====================================================================

_TOKEN_RE = re.compile(r"[a-zA-Z0-9_]+|[\u4e00-\u9fff]+")
_SENT_SPLIT_RE = re.compile(r"(?<=[。！？!?；;\n])")
_PAGE_MARKER_RE = re.compile(r"\[\[PAGE:(\d+)\]\]")


def tokenize(text: str) -> List[str]:
    """中英文分词：英文按 `[a-zA-Z0-9_]+`，中文按字 + 2~4 字 n-gram。

    中文采用 n-gram 是为了让 BM25 能精确命中「正则表达式」「检索」等多字术语，
    同时保留单字以兜底单字查询。
    """
    tokens: List[str] = []
    for raw in _TOKEN_RE.findall(text.lower()):
        if re.fullmatch(r"[\u4e00-\u9fff]+", raw):
            tokens.extend(raw)  # 单字
            for size in range(2, min(5, len(raw)) + 1):
                tokens.extend(
                    raw[i:i + size] for i in range(len(raw) - size + 1)
                )
        else:
            tokens.append(raw)
    return tokens


def bm25_scores(query: str, documents: List[str],
                k1: float = 1.5, b: float = 0.75) -> List[float]:
    """对每篇文档计算查询的 BM25 分数（纯 Python，不依赖 Milvus）。"""
    q_tokens = tokenize(query)
    doc_tokens = [tokenize(d) for d in documents]
    n_docs = len(doc_tokens)
    if n_docs == 0:
        return []

    df: Dict[str, int] = {}
    for dt in doc_tokens:
        for term in set(dt):
            df[term] = df.get(term, 0) + 1
    avgdl = sum(len(dt) for dt in doc_tokens) / n_docs

    scores: List[float] = []
    for dt in doc_tokens:
        tf: Dict[str, int] = {}
        for term in dt:
            tf[term] = tf.get(term, 0) + 1
        dl = len(dt)
        score = 0.0
        for term in q_tokens:
            if term not in tf:
                continue
            n = df.get(term, 0)
            idf = math.log((n_docs - n + 0.5) / (n + 0.5) + 1)
            score += idf * (tf[term] * (k1 + 1)) / (
                tf[term] + k1 * (1 - b + b * dl / (avgdl or 1)))
        scores.append(score)
    return scores


def chunk_content_hash(text: str) -> str:
    """基于内容（SHA-256）生成分块标识，相同文本得到相同 id。"""
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def split_document_pages(text: str) -> List[tuple]:
    """解析 `[[PAGE:n]]` 页码标记，返回 [(page, page_text), ...]。

    PDF 导入时会插入页码标记，用于检索结果的来源定位。
    无标记时返回空列表（由调用方决定回退策略）。
    """
    parts = _PAGE_MARKER_RE.split(text)
    pages: List[tuple] = []
    # parts 结构：[pre_text, page1, text1, page2, text2, ...]
    i = 1
    while i < len(parts):
        page = int(parts[i])
        content = parts[i + 1].strip() if i + 1 < len(parts) else ""
        if content:
            pages.append((page, content))
        i += 2
    return pages


def _cosine(a: List[float], b: List[float]) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a))
    nb = math.sqrt(sum(y * y for y in b))
    if na == 0 or nb == 0:
        return 0.0
    return dot / (na * nb)


def _percentile(values: List[float], pct: float) -> float:
    """线性插值百分位，避免对 numpy 的硬依赖。"""
    if not values:
        return 0.0
    s = sorted(values)
    if len(s) == 1:
        return float(s[0])
    k = (len(s) - 1) * (pct / 100.0)
    f = int(k)
    c = k - f
    if f + 1 < len(s):
        return s[f] + (s[f + 1] - s[f]) * c
    return float(s[f])


def semantic_breakpoints(embeddings: List[List[float]],
                         percentile: float = 25) -> List[int]:
    """基于相邻句向量余弦相似度识别主题边界。

    相似度低于 `percentile` 分位的相邻位置视为语义跳变点（断点）。
    返回断点在 sentences 中的起始下标。
    """
    if len(embeddings) < 2:
        return []
    sims = [_cosine(embeddings[i], embeddings[i + 1])
            for i in range(len(embeddings) - 1)]
    if not sims:
        return []
    threshold = _percentile(sims, percentile)
    points: List[int] = []
    for i, sim in enumerate(sims):
        if sim < threshold:
            points.append(i + 1)
    return points


def merge_semantic_sentences(sentences: List[str], breakpoints: List[int],
                             target_chars: int = 400) -> List[str]:
    """按断点把句子分组合并成块，每组内句子直接拼接。"""
    if not sentences:
        return []
    bps = sorted(set(breakpoints))
    bps_with_end = list(bps) + [len(sentences)]
    chunks: List[str] = []
    start = 0
    for bp in bps_with_end:
        if bp <= start:
            continue
        group = sentences[start:bp]
        if group:
            chunks.append("".join(group))
        start = bp
    return chunks


def reciprocal_rank_fusion(rankings: List[tuple], limit: int = 10,
                           k: int = 60) -> List[str]:
    """倒数排名融合。rankings: [(ids_list, weight), ...]。

    score(id) = Σ weight / (k + rank)。用于融合 Dense / Sparse 两路检索结果。
    """
    scores: Dict[str, float] = {}
    for ids, weight in rankings:
        for rank, id_ in enumerate(ids):
            scores[id_] = scores.get(id_, 0.0) + weight / (k + rank)
    ordered = sorted(scores, key=lambda x: (-scores[x], x))
    return ordered[:limit]


def _split_sentences(text: str) -> List[str]:
    text = text.strip()
    if not text:
        return []
    return [s.strip() for s in _SENT_SPLIT_RE.split(text) if s.strip()]


def _rule_chunk(text: str, max_chars: int = 1200,
                overlap: int = 100) -> List[str]:
    """带重叠窗口的规则分块（语义分块异常时的回退策略）。"""
    text = text.strip()
    if not text:
        return []
    if len(text) <= max_chars:
        return [text]
    sents = _split_sentences(text)
    chunks: List[str] = []
    buf = ""
    for s in sents:
        if len(buf) + len(s) <= max_chars:
            buf += s
        else:
            if buf:
                chunks.append(buf)
            buf = (buf[-overlap:] if len(buf) > overlap else "") + s
    if buf:
        chunks.append(buf)
    return chunks


def _semantic_chunk(text: str, target_chars: int = 400,
                    max_chars: int = 1200) -> List[str]:
    """语义分块：句向量识别主题边界 → 合并；异常时回退规则分块。"""
    sentences = _split_sentences(text)
    if len(sentences) <= 1:
        return [text] if text.strip() else []
    try:
        embeddings = _engine().encode_documents(sentences)
        breakpoints = semantic_breakpoints(embeddings)
        chunks = merge_semantic_sentences(sentences, breakpoints, target_chars)
        # 过长块按规则再切
        out: List[str] = []
        for c in chunks:
            if len(c) <= max_chars:
                out.append(c)
            else:
                out.extend(_rule_chunk(c, max_chars))
        return out
    except Exception as e:  # noqa: BLE001 - 语义分块失败需兜底
        print(f"[semantic_chunk] 语义分块失败，回退规则分块: {e}")
        return _rule_chunk(text, max_chars)


async def rewrite_retrieval_query(req: RagRetrieveRequest) -> str:
    """结合对话历史改写检索查询。

    - 无历史（首问）：直接返回原查询，不调用 LLM。
    - 有历史（追问）：用 LLM 把指代/省略类追问改写为可独立检索的完整查询。
    - 改写失败：回退为原查询，保证检索不中断。
    """
    if not req.history:
        return req.query
    history_text = "\n".join(
        f"{m.get('role', 'user')}: {m.get('content', '')}" for m in req.history
    )
    system = (
        "你是查询改写助手。根据对话历史，将用户的追问改写为一个可独立检索的"
        "完整查询，补全省略的主语或指代对象。只返回改写后的查询本身，"
        "不要解释、不要加引号、不要任何前缀。"
    )
    user = (f"对话历史：\n{history_text}\n\n"
            f"用户追问：{req.query}\n\n改写后的独立查询：")
    try:
        return await call_llm(
            req.rewrite_provider or "deepseek",
            system, user,
            req.rewrite_api_key, req.rewrite_base_url, req.rewrite_model,
            temperature=0.2,
        )
    except Exception as e:  # noqa: BLE001 - 改写失败必须兜底
        print(f"[rewrite_retrieval_query] 改写失败，回退原查询: {e}")
        return req.query


# ====================================================================
# 云端 LLM 调用（摘要 / 润色 / 查询改写共用）
# ====================================================================
async def call_llm(provider: str, system_prompt: str, user_prompt: str,
                   req_key: Optional[str], req_url: Optional[str],
                   req_model: Optional[str], temperature: float = 0.3) -> str:
    """调用云端 OpenAI 兼容接口，返回模型文本。"""
    if provider not in DEFAULT_PROVIDERS:
        provider = "deepseek"
    cfg = DEFAULT_PROVIDERS[provider]

    api_key = req_key or cfg["api_key"]
    base_url = (req_url or cfg["base_url"]).rstrip("/")
    model = req_model or cfg["model"]

    if not api_key:
        raise HTTPException(
            status_code=400,
            detail=f"服务商 [{provider}] 未配置 API Key（请在 .env 或请求中提供）")

    url = f"{base_url}/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "temperature": temperature,
    }

    async with httpx.AsyncClient(timeout=90.0) as client:
        try:
            resp = await client.post(url, headers=headers, json=payload)
        except httpx.ConnectError:
            raise HTTPException(status_code=502,
                                detail="无法连接模型服务（BaseURL 错误或网络不通）")
        if resp.status_code != 200:
            raise HTTPException(status_code=resp.status_code,
                                detail=f"模型服务返回错误：{resp.text[:500]}")
        data = resp.json()
        choices = data.get("choices") or []
        if not choices:
            raise HTTPException(status_code=502,
                                detail=f"模型返回异常（choices 为空）：{str(data)[:300]}")
        msg = choices[0].get("message", {})
        content = msg.get("content", "")
        if not content:
            raise HTTPException(status_code=502,
                                detail="模型返回内容为空")
        return content.strip()


async def call_llm_stream(provider: str, system_prompt: str, user_prompt: str,
                          req_key: Optional[str], req_url: Optional[str],
                          req_model: Optional[str], temperature: float = 0.3):
    """流式调用云端 OpenAI 兼容接口，逐块 yield 文本片段。"""
    import asyncio
    if provider not in DEFAULT_PROVIDERS:
        provider = "deepseek"
    cfg = DEFAULT_PROVIDERS[provider]

    api_key = req_key or cfg["api_key"]
    base_url = (req_url or cfg["base_url"]).rstrip("/")
    model = req_model or cfg["model"]

    if not api_key:
        raise HTTPException(
            status_code=400,
            detail=f"服务商 [{provider}] 未配置 API Key（请在 .env 或请求中提供）")

    url = f"{base_url}/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "temperature": temperature,
        "stream": True,
    }
    print(f"[call_llm_stream] url={url}  model={model}  provider={provider}")

    yielded = False
    last_error = None
    for attempt in range(3):
        try:
            async with httpx.AsyncClient(timeout=120.0) as client:
                async with client.stream("POST", url, headers=headers, json=payload) as resp:
                    if resp.status_code != 200:
                        body = await resp.aread()
                        raise HTTPException(status_code=resp.status_code,
                                            detail=f"模型服务返回错误：{body.decode()[:500]}")
                    async for line in resp.aiter_lines():
                        if line.startswith("data: "):
                            chunk = line[6:]
                            if chunk.strip() == "[DONE]":
                                break
                            try:
                                obj = json.loads(chunk)
                                delta = (obj.get("choices") or [{}])[0].get("delta", {})
                                text = delta.get("content", "")
                                if text:
                                    yielded = True
                                    yield text
                            except (json.JSONDecodeError, IndexError, KeyError):
                                continue
            return
        except (httpx.RemoteProtocolError, httpx.ConnectError,
                httpx.ReadTimeout, httpx.ConnectTimeout) as e:
            last_error = e
            if yielded:
                raise HTTPException(status_code=502,
                                    detail=f"流式传输中断：{e}")
            if attempt < 2:
                print(f"[call_llm_stream] 第 {attempt+1} 次失败({e})，{attempt+1}s 后重试…")
                await asyncio.sleep(attempt + 1)
            continue
    raise HTTPException(status_code=502,
                        detail=f"多次重试后仍失败：{last_error}")


# ---- 摘要 / 学术写作辅助 ----
@app.post("/api/summary")
async def summary(req: TextRequest):
    level_map = {
        "short": "一句话的",
        "medium": "一段简洁的",
        "detailed": "详细的、保留关键数据与细节的",
    }
    level = level_map.get(req.detail, "一段简洁的")
    system = (f"你是一个专业的文本摘要助手。请为用户输入的笔记文本生成{level}摘要，"
              f"保留核心观点、关键信息和数据，不要编造原文没有的内容。")
    result = await call_llm(req.provider, system, req.text,
                            req.api_key, req.base_url, req.model)
    return {"summary": result}


@app.post("/api/polish")
async def polish(req: TextRequest):
    system = (
        "你是学术写作辅助专家。在不改变事实、论点和引用关系的前提下，"
        "优化表达、逻辑衔接、术语一致性和语言简洁度。\n\n"
        "润色原则（按优先级排序）：\n"
        "1. 修正语法错误、错别字、标点误用、病句；\n"
        "2. 删去重复啰嗦的表述，提升语言简洁度；\n"
        "3. 调整语序与衔接，让逻辑更连贯；\n"
        "4. 统一术语，仅当原文用词明显不当时替换为更精准的表达。\n\n"
        "硬性约束：\n"
        "- 必须保留原文的核心观点、事实信息、引用关系和语气风格；\n"
        "- 若某句话已通顺得体则原样保留，不要强行改写；\n"
        "- 仅返回润色后的完整正文，不要任何解释、前缀或元评论。"
    )
    result = await call_llm(req.provider, system, req.text,
                            req.api_key, req.base_url, req.model,
                            temperature=0.5)
    return {"polished": result}


# ---- 流式接口 ----
@app.post("/api/summary/stream")
async def summary_stream(req: TextRequest):
    level_map = {
        "short": "一句话的",
        "medium": "一段简洁的",
        "detailed": "详细的、保留关键数据与细节的",
    }
    level = level_map.get(req.detail, "一段简洁的")
    system = (f"你是一个专业的文本摘要助手。请为用户输入的笔记文本生成{level}摘要，"
              f"保留核心观点、关键信息和数据，不要编造原文没有的内容。")

    async def gen():
        async for chunk in call_llm_stream(req.provider, system, req.text,
                                           req.api_key, req.base_url, req.model):
            yield chunk
    return StreamingResponse(gen(), media_type="text/plain; charset=utf-8")


@app.post("/api/polish/stream")
async def polish_stream(req: TextRequest):
    system = (
        "你是学术写作辅助专家。在不改变事实、论点和引用关系的前提下，"
        "优化表达、逻辑衔接、术语一致性和语言简洁度。"
        "仅返回润色后的完整正文，不要任何解释或前缀。"
    )
    async def gen():
        async for chunk in call_llm_stream(req.provider, system, req.text,
                                           req.api_key, req.base_url, req.model,
                                           temperature=0.5):
            yield chunk
    return StreamingResponse(gen(), media_type="text/plain; charset=utf-8")


# ====================================================================
# RAG 端点：本地 Jina + Milvus Lite 混合检索
# ====================================================================
@app.post("/api/rag/index")
async def rag_index(req: RagIndexRequest):
    """增量建立或更新文献索引：页码解析 → 语义分块 → 内容哈希 → 入库。

    通过 SHA-256 内容哈希识别分块，未变化分块复用已缓存向量，
    仅对新增/变化分块调用本地 Embedding。
    """
    text = req.text.strip()
    if not text:
        return {"indexed": 0}
    if len(text) > 1_000_000:
        raise HTTPException(status_code=413,
                            detail="单次索引请求上限为 100 万字符")

    pages = split_document_pages(text)
    if not pages:
        pages = [(0, text)]

    records: List[dict] = []
    for page_no, page_text in pages:
        chunks = _semantic_chunk(page_text)
        for idx, chunk in enumerate(chunks):
            chunk = chunk.strip()
            if not chunk:
                continue
            records.append({
                "id": chunk_content_hash(f"{req.note_id}:{page_no}:{idx}:{chunk}"),
                "text": chunk,
                "title": req.title or "未命名文献",
                "source": req.source,
                "page": page_no,
                "idx": idx,
                "hash": chunk_content_hash(chunk),
            })

    if not records:
        return {"indexed": 0}

    print(f"[rag_index] note_id={req.note_id}  原文{len(text)}字  "
          f"分块{len(records)}个")
    result = _engine().replace_note(req.note_id, records)
    return {"indexed": result.get("indexed", len(records))}


@app.post("/api/rag/retrieve")
async def rag_retrieve(req: RagRetrieveRequest):
    """混合检索：查询改写 → Dense+Sparse 混合检索 → 返回片段与来源元数据。

    document_ids 为空/None 时检索全部文献；非空时限定在这些文献内。
    返回 {"chunks": [str], "sources": [{title, page, source}]}。
    """
    rewritten = await rewrite_retrieval_query(req)
    document_ids = req.document_ids if req.document_ids else None
    hits = _engine().search(rewritten, req.top_k, document_ids)

    chunks = [hit.get("text", "") for hit in hits]
    sources = [{
        "title": hit.get("title", "未命名文献"),
        "page": hit.get("page", 0),
        "source": hit.get("source", ""),
    } for hit in hits]

    print(f"[rag_retrieve] query='{req.query}'  rewritten='{rewritten}'  "
          f"document_ids={document_ids}  top_k={req.top_k}  命中 {len(hits)} 个片段")
    for i, (c, s) in enumerate(zip(chunks, sources)):
        print(f"  [{i+1}] ({len(c)}字, 第{s['page']}页, {s['title']}) {c[:80]}...")

    return {"chunks": chunks, "sources": sources}


@app.post("/api/rag/delete")
async def rag_delete(req: RagDeleteRequest):
    """删除指定文献的全部向量索引。"""
    _engine().delete_note(req.note_id)
    return {"deleted": True}


@app.get("/api/rag/stats")
async def rag_stats():
    """返回本地 RAG 引擎统计信息（文献数、分块数等）。"""
    return _engine().stats()


# ---- 文档导入：从 PDF / Word / TXT / Markdown 提取纯文本 ----
@app.post("/api/import")
async def import_document(file: UploadFile = File(...)):
    """上传文档，提取纯文本返回。PDF 保留页码标记 [[PAGE:n]]。"""
    import io

    filename = file.filename or "未命名"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    title = filename.rsplit(".", 1)[0] if "." in filename else filename

    raw = await file.read()
    if len(raw) > 50 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="文件过大（超过 50MB），请拆分后导入")

    try:
        if ext == "pdf":
            import pdfplumber
            pages = []
            with pdfplumber.open(io.BytesIO(raw)) as doc:
                for pg_no, pg in enumerate(doc.pages, start=1):
                    t = pg.extract_text() or ""
                    pages.append(f"[[PAGE:{pg_no}]]\n{t}")
            text = "\n\n".join(pages).strip()
            if not text:
                raise HTTPException(status_code=422,
                                    detail="未能从 PDF 中提取到文字，可能是扫描版 PDF（需 OCR 支持）")

        elif ext == "docx":
            from docx import Document
            doc = Document(io.BytesIO(raw))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            text = "\n\n".join(paragraphs)
            if not text:
                raise HTTPException(status_code=422, detail="Word 文档内容为空")

        elif ext in ("txt", "md", "markdown"):
            # 纯文本 / Markdown：直接读取（尝试常见编码）
            for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
                try:
                    text = raw.decode(enc).strip()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                text = raw.decode("utf-8", errors="ignore").strip()
            if not text:
                raise HTTPException(status_code=422, detail="文件内容为空")

        else:
            raise HTTPException(status_code=400,
                                detail=f"不支持的文件类型 .{ext}，"
                                      f"仅支持 .pdf / .docx / .txt / .md")

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文档解析失败：{e}")

    return {"title": title, "content": text}


@app.get("/health")
async def health():
    return {"status": "ok"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
